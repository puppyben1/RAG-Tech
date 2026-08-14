import csv
import json
from pathlib import Path

from docx import Document

from jinrong.qa_review_import import EVIDENCE_CONTAINMENT_MIN, _evidence_containment, import_qa_reviews


FIELDS = [
    "id",
    "selected_doc_id",
    "page_no",
    "article_no",
    "sheet_name",
    "cell_ref",
    "review_notes",
]


def _write_jsonl(path: Path, rows: list[dict]) -> None:
    path.write_text("".join(json.dumps(row, ensure_ascii=False) + "\n" for row in rows), encoding="utf-8")


def _write_reviews(path: Path, rows: list[dict[str, str]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=FIELDS)
        writer.writeheader()
        writer.writerows(rows)


def _candidate(candidate_ids: list[str]) -> dict:
    return {
        "id": "Q101",
        "migration_status": "document_matched_locator_missing",
        "source_type": "word",
        "original_evidence": "机构应当依法经营。",
        "candidate_doc_ids": candidate_ids,
    }


def test_blank_review_remains_fail_closed(tmp_path: Path) -> None:
    review = tmp_path / "review.csv"
    candidates = tmp_path / "candidates.jsonl"
    manifest = tmp_path / "manifest.jsonl"
    output = tmp_path / "reviewed.jsonl"
    _write_reviews(review, [{"id": "Q101", "selected_doc_id": ""}])
    _write_jsonl(candidates, [_candidate(["doc-1"])])
    _write_jsonl(manifest, [{"doc_id": "doc-1", "file_name": "one.docx"}])

    result = import_qa_reviews(review, candidates, output, manifest, tmp_path)

    assert result["status"] == "review_required"
    assert result["not_reviewed"] == 1
    assert output.read_text(encoding="utf-8") == ""


def test_resolved_review_cannot_select_unrelated_manifest_document(tmp_path: Path) -> None:
    review = tmp_path / "review.csv"
    candidates = tmp_path / "candidates.jsonl"
    manifest = tmp_path / "manifest.jsonl"
    output = tmp_path / "reviewed.jsonl"
    _write_reviews(review, [{"id": "Q101", "selected_doc_id": "doc-2", "article_no": "第一条"}])
    _write_jsonl(candidates, [_candidate(["doc-1"])])
    _write_jsonl(
        manifest,
        [
            {"doc_id": "doc-1", "file_name": "one.docx"},
            {"doc_id": "doc-2", "file_name": "two.docx"},
        ],
    )

    result = import_qa_reviews(review, candidates, output, manifest, tmp_path)

    assert result["status"] == "validation_failed"
    assert result["errors"][0]["error"] == "selected_doc_not_in_candidates"


def test_docx_article_must_exist_in_document(tmp_path: Path) -> None:
    data_dir = tmp_path / "wendang" / "data"
    data_dir.mkdir(parents=True)
    document = Document()
    document.add_paragraph("第一条 机构应当依法经营。")
    document.save(data_dir / "one.docx")
    review = tmp_path / "review.csv"
    candidates = tmp_path / "candidates.jsonl"
    manifest = tmp_path / "manifest.jsonl"
    output = tmp_path / "reviewed.jsonl"
    _write_reviews(review, [{"id": "Q101", "selected_doc_id": "doc-1", "article_no": "第二条"}])
    _write_jsonl(candidates, [_candidate(["doc-1"])])
    _write_jsonl(manifest, [{"doc_id": "doc-1", "file_name": "one.docx"}])

    result = import_qa_reviews(review, candidates, output, manifest, tmp_path)

    assert result["status"] == "validation_failed"
    assert result["errors"][0]["error"] == "locator_validation_failed"


def test_valid_docx_review_is_imported(tmp_path: Path) -> None:
    data_dir = tmp_path / "wendang" / "data"
    data_dir.mkdir(parents=True)
    document = Document()
    document.add_paragraph("第一条 机构应当依法经营。")
    document.save(data_dir / "one.docx")
    review = tmp_path / "review.csv"
    candidates = tmp_path / "candidates.jsonl"
    manifest = tmp_path / "manifest.jsonl"
    output = tmp_path / "reviewed.jsonl"
    _write_reviews(review, [{"id": "Q101", "selected_doc_id": "doc-1", "article_no": "第一条"}])
    _write_jsonl(candidates, [_candidate(["doc-1"])])
    _write_jsonl(manifest, [{"doc_id": "doc-1", "file_name": "one.docx"}])

    result = import_qa_reviews(review, candidates, output, manifest, tmp_path)
    row = json.loads(output.read_text(encoding="utf-8"))

    assert result["status"] == "passed"
    assert row["migration_status"] == "reviewed_and_ready"
    assert row["gold_evidence"] == [{"doc_id": "doc-1", "article_no": "第一条"}]

def test_evidence_containment_exact_match_is_one() -> None:
    assert _evidence_containment("终极利率暂定为4.5%。", "终极利率暂定为 4.5%。") == 1.0


def test_evidence_containment_accepts_paraphrased_evidence() -> None:
    evidence = "计算现金流现值所采用的折现率曲线由基础利率曲线加综合溢价形成。"
    page = "根据《保险公司偿付能力监管规则第3号》规定，计算现金流现值所采用的折现率曲线由基础利率曲线加综合溢价形成，具体计算方法如下："
    assert _evidence_containment(evidence, page) >= EVIDENCE_CONTAINMENT_MIN


def test_evidence_containment_rejects_unrelated_text() -> None:
    evidence = "寿险合同负债评估中计算现金流现值所采用的折现率曲线"
    page = "银行函证工作操作指引用于进一步明确和细化银行函证工作中的具体事项"
    assert _evidence_containment(evidence, page) < EVIDENCE_CONTAINMENT_MIN
