from __future__ import annotations

import re
from functools import lru_cache
from pathlib import Path

import pdfplumber
from docx import Document

from .utils import norm_text


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n[page {i}]\n{text}")
    return "\n".join(parts)


def extract_doc_binary_fallback(path: Path) -> str:
    """Best-effort old .doc extraction; reliable .doc support needs external converters."""
    data = path.read_bytes()
    candidates: list[str] = []
    for encoding in ("utf-16le", "gb18030", "utf-8"):
        decoded = data.decode(encoding, errors="ignore")
        snippets = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：《》（）“”%、\-—]{8,}", decoded)
        candidates.extend(snippets)
    seen: set[str] = set()
    kept: list[str] = []
    for text in candidates:
        key = norm_text(text)
        if key and key not in seen:
            seen.add(key)
            kept.append(text)
    return "\n".join(kept)


@lru_cache(maxsize=256)
def extract_text(path_str: str) -> str:
    path = Path(path_str)
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".doc":
        return extract_doc_binary_fallback(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def split_sentences(text: str) -> list[str]:
    raw = re.split(r"(?<=[。！？；])\s*", text)
    return [s.strip() for s in raw if s.strip()]

