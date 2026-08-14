"""Import and validate human-reviewed QA migration results."""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
import pdfplumber
from docx import Document

from .config import MANIFEST_PATH, PROJECT_ROOT
from .utils import norm_text, read_jsonl, write_jsonl

# Minimum fraction of evidence character bigrams that must be present on a
# cited PDF page for the locator to pass. QA evidence strings are frequently
# paraphrases of the source text, so an exact substring match is too strict;
# bigram containment keeps the check fail-closed while tolerating minor
# rewording.
EVIDENCE_CONTAINMENT_MIN = 0.5


def import_qa_reviews(
    review_csv_path: Path,
    candidate_jsonl_path: Path,
    output_jsonl_path: Path,
    manifest_path: Path = MANIFEST_PATH,
    project_root: Path = PROJECT_ROOT,
) -> dict[str, Any]:
    """
    Import human-reviewed QA migration results from CSV.

    Validates:
    - selected_doc_id must come from candidates or exist in manifest
    - Word docs must have article_no or page_no
    - PDF docs must have page_no
    - Excel docs must have sheet_name and cell_ref
    - Files must exist and locators must be valid

    Returns validation report and writes merged results to output.
    """
    input_paths = {review_csv_path.resolve(), candidate_jsonl_path.resolve(), manifest_path.resolve()}
    if output_jsonl_path.resolve() in input_paths:
        raise ValueError("output_jsonl_path must differ from every input path")

    if not review_csv_path.is_file():
        raise ValueError(f"review CSV not found: {review_csv_path}")

    if not candidate_jsonl_path.is_file():
        raise ValueError(f"candidate JSONL not found: {candidate_jsonl_path}")

    # Load data
    review_rows = _load_review_csv(review_csv_path)
    review_ids = [row.get("id", "").strip() for row in review_rows]
    duplicate_ids = sorted({qa_id for qa_id in review_ids if qa_id and review_ids.count(qa_id) > 1})
    if duplicate_ids:
        raise ValueError(f"duplicate QA IDs in review CSV: {', '.join(duplicate_ids)}")
    candidates = {row["id"]: row for row in read_jsonl(candidate_jsonl_path)}
    manifest = read_jsonl(manifest_path)
    manifest_by_id = {row["doc_id"]: row for row in manifest}

    # Validate and merge
    passed: list[dict[str, Any]] = []
    failed: list[dict[str, Any]] = []
    validation_errors: list[dict[str, Any]] = []

    for review in review_rows:
        qa_id = review["id"]
        if qa_id not in candidates:
            validation_errors.append({
                "id": qa_id,
                "error": "not_found_in_candidates",
                "message": f"QA ID {qa_id} not found in candidate JSONL"
            })
            continue

        candidate = candidates[qa_id]
        selected_doc_id = review.get("selected_doc_id", "").strip()

        # Skip if not reviewed (no selected_doc_id)
        if not selected_doc_id:
            failed.append({
                "id": qa_id,
                "reason": "not_reviewed",
                "original_status": candidate.get("migration_status")
            })
            continue

        # Resolved and ambiguous rows must stay within their discovered candidate set.
        # Unresolved rows have no candidates, so a reviewer may select any manifest document.
        candidate_ids = [str(value) for value in candidate.get("candidate_doc_ids", [])]
        if candidate_ids and selected_doc_id not in candidate_ids:
            validation_errors.append({
                "id": qa_id,
                "error": "selected_doc_not_in_candidates",
                "selected_doc_id": selected_doc_id,
                "candidate_doc_ids": candidate_ids,
            })
            failed.append({
                "id": qa_id,
                "reason": "selected_doc_not_in_candidates",
                "selected_doc_id": selected_doc_id,
            })
            continue
        if selected_doc_id not in manifest_by_id:
            validation_errors.append({
                "id": qa_id,
                "error": "doc_id_not_in_manifest",
                "selected_doc_id": selected_doc_id,
            })
            failed.append({
                "id": qa_id,
                "reason": "doc_id_not_in_manifest",
                "selected_doc_id": selected_doc_id,
            })
            continue

        doc_record = manifest_by_id[selected_doc_id]
        source_type = candidate.get("source_type", "").lower()

        # Build locators based on source type
        locators = {}
        locator_errors = []

        if source_type == "excel":
            sheet_name = review.get("sheet_name", "").strip()
            cell_ref = review.get("cell_ref", "").strip()

            if not sheet_name or not cell_ref:
                locator_errors.append("Excel requires both sheet_name and cell_ref")
            else:
                locators["sheet_name"] = sheet_name
                locators["cell_ref"] = cell_ref

        elif source_type == "word":
            article_no = review.get("article_no", "").strip()
            page_no = review.get("page_no", "").strip()

            if not article_no and not page_no:
                locator_errors.append("Word requires at least article_no or page_no")
            else:
                if article_no:
                    locators["article_no"] = article_no
                if page_no:
                    locators["page_no"] = _parse_page_number(page_no)

        elif source_type == "pdf":
            page_no = review.get("page_no", "").strip()

            if not page_no:
                locator_errors.append("PDF requires page_no")
            else:
                locators["page_no"] = _parse_page_number(page_no)

        if any(value is None for value in locators.values()):
            locator_errors.append("page_no must be a positive integer")

        if locator_errors:
            validation_errors.append({
                "id": qa_id,
                "error": "missing_locators",
                "source_type": source_type,
                "errors": locator_errors
            })
            failed.append({
                "id": qa_id,
                "reason": "missing_locators",
                "errors": locator_errors
            })
            continue

        # Validate file exists
        file_name = doc_record.get("file_name")
        if not file_name:
            validation_errors.append({
                "id": qa_id,
                "error": "no_file_name",
                "selected_doc_id": selected_doc_id
            })
            failed.append({
                "id": qa_id,
                "reason": "no_file_name",
                "selected_doc_id": selected_doc_id
            })
            continue

        file_path = project_root / "wendang" / "data" / file_name
        if not file_path.is_file():
            validation_errors.append({
                "id": qa_id,
                "error": "file_not_found",
                "file_path": str(file_path)
            })
            failed.append({
                "id": qa_id,
                "reason": "file_not_found",
                "file_path": str(file_path)
            })
            continue

        # Validate locators against actual file
        validation_result = _validate_locators_in_file(
            file_path, source_type, locators, candidate.get("original_evidence", "")
        )

        if not validation_result["valid"]:
            validation_errors.append({
                "id": qa_id,
                "error": "locator_validation_failed",
                "file_path": str(file_path),
                "locators": locators,
                "validation_errors": validation_result.get("errors", [])
            })
            failed.append({
                "id": qa_id,
                "reason": "locator_validation_failed",
                "validation_errors": validation_result.get("errors", [])
            })
            continue

        # Build reviewed candidate
        reviewed = dict(candidate)
        reviewed["migration_status"] = "reviewed_and_ready"
        reviewed["doc_id"] = selected_doc_id
        reviewed["selected_doc_id"] = selected_doc_id
        reviewed["expected_doc_ids"] = [selected_doc_id]

        # Build gold_evidence
        gold_evidence = [{"doc_id": selected_doc_id, **locators}]
        reviewed["gold_evidence"] = gold_evidence

        # Update local_path
        reviewed["local_path"] = f"wendang/data/{file_name}"

        # Add review metadata
        review_notes = review.get("review_notes", "").strip()
        if review_notes:
            reviewed["review_notes"] = review_notes

        passed.append(reviewed)

    # Write output
    output_jsonl_path.parent.mkdir(parents=True, exist_ok=True)
    write_jsonl(output_jsonl_path, passed)

    # Build summary
    status = "validation_failed" if validation_errors else ("review_required" if failed or not review_rows else "passed")
    return {
        "status": status,
        "review_csv": str(review_csv_path),
        "candidate_jsonl": str(candidate_jsonl_path),
        "output_jsonl": str(output_jsonl_path),
        "total_reviews": len(review_rows),
        "passed": len(passed),
        "failed": len(failed),
        "not_reviewed": len([f for f in failed if f.get("reason") == "not_reviewed"]),
        "validation_errors": len(validation_errors),
        "errors": validation_errors if validation_errors else None,
        "failed_items": failed if failed else None,
    }


def _load_review_csv(path: Path) -> list[dict[str, str]]:
    """Load review CSV with UTF-8 BOM support."""
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        return list(reader)


def _parse_page_number(page_str: str) -> int | None:
    """Parse a strictly positive integer PDF page number."""
    value = page_str.strip()
    return int(value) if re.fullmatch(r"[1-9]\d*", value) else None


def _validate_locators_in_file(
    file_path: Path,
    source_type: str,
    locators: dict[str, Any],
    original_evidence: str,
) -> dict[str, Any]:
    """Validate that locators are valid for the given file."""
    errors = []

    try:
        if source_type == "excel":
            result = _validate_excel_locators(file_path, locators)
            return result

        elif source_type == "pdf":
            result = _validate_pdf_locators(file_path, locators, original_evidence)
            return result

        elif source_type == "word":
            return _validate_word_locators(file_path, locators)

        else:
            errors.append(f"Unsupported source_type: {source_type}")
            return {"valid": False, "errors": errors}

    except Exception as e:
        errors.append(f"Validation exception: {str(e)}")
        return {"valid": False, "errors": errors}


def _validate_excel_locators(file_path: Path, locators: dict[str, Any]) -> dict[str, Any]:
    """Validate Excel sheet and cell reference."""
    errors = []
    sheet_name = locators.get("sheet_name")
    cell_ref = locators.get("cell_ref")

    if not sheet_name or not cell_ref:
        errors.append("Missing sheet_name or cell_ref")
        return {"valid": False, "errors": errors}

    # Validate cell reference format
    if not re.match(r"^[A-Z]+\d+$", cell_ref, re.IGNORECASE):
        errors.append(f"Invalid cell reference format: {cell_ref}")
        return {"valid": False, "errors": errors}

    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        try:
            # Check if sheet exists
            actual_sheet_name = sheet_name
            if actual_sheet_name not in workbook.sheetnames:
                matched = [s for s in workbook.sheetnames if norm_text(s) == norm_text(sheet_name)]
                if len(matched) != 1:
                    errors.append(f"Sheet '{sheet_name}' not found. Available: {workbook.sheetnames}")
                    return {"valid": False, "errors": errors}
                actual_sheet_name = matched[0]

            # Try to access the cell
            sheet = workbook[actual_sheet_name]
            cell_value = sheet[cell_ref].value

            if cell_value is None or (isinstance(cell_value, str) and not cell_value.strip()):
                errors.append(f"Cell '{actual_sheet_name}!{cell_ref}' is empty")
                return {"valid": False, "errors": errors}
            return {"valid": True}

        finally:
            workbook.close()

    except Exception as e:
        errors.append(f"Excel validation error: {str(e)}")
        return {"valid": False, "errors": errors}


def _validate_pdf_locators(
    file_path: Path,
    locators: dict[str, Any],
    original_evidence: str,
) -> dict[str, Any]:
    """Validate PDF page number."""
    errors = []
    page_no = locators.get("page_no")

    if page_no is None:
        errors.append("Missing page_no")
        return {"valid": False, "errors": errors}

    try:
        page_no = int(page_no)
    except (ValueError, TypeError):
        errors.append(f"Invalid page number: {page_no}")
        return {"valid": False, "errors": errors}

    if page_no < 1:
        errors.append(f"Page number must be >= 1: {page_no}")
        return {"valid": False, "errors": errors}

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)

            if page_no > page_count:
                errors.append(f"Page {page_no} exceeds PDF page count {page_count}")
                return {"valid": False, "errors": errors}

            evidence = norm_text(original_evidence)
            page_text = norm_text(pdf.pages[page_no - 1].extract_text() or "")
            if evidence and evidence not in page_text:
                containment = _evidence_containment(evidence, page_text)
                if containment < EVIDENCE_CONTAINMENT_MIN:
                    best_page, best_score = _best_evidence_page(pdf, evidence)
                    errors.append(
                        f"Original evidence similarity {containment:.2f} is below "
                        f"{EVIDENCE_CONTAINMENT_MIN} on PDF page {page_no} "
                        f"(best match: page {best_page}, similarity {best_score:.2f})"
                    )
                    return {"valid": False, "errors": errors}

        return {"valid": True}

    except Exception as e:
        errors.append(f"PDF validation error: {str(e)}")
        return {"valid": False, "errors": errors}




def _bigrams(value: str) -> set[str]:
    """Character bigrams used for fuzzy Chinese text matching."""
    return {value[i:i + 2] for i in range(len(value) - 1)}


def _evidence_containment(evidence: str, text: str) -> float:
    """Fraction of evidence character bigrams contained in text.

    Both inputs are whitespace-normalized first. Returns 1.0 when the
    normalized evidence is a substring of the normalized text.
    """
    evidence_norm = norm_text(evidence)
    text_norm = norm_text(text)
    if not evidence_norm or not text_norm:
        return 0.0
    if evidence_norm in text_norm:
        return 1.0
    evidence_bigrams = _bigrams(evidence_norm)
    if not evidence_bigrams:
        return 0.0
    return len(evidence_bigrams & _bigrams(text_norm)) / len(evidence_bigrams)


def _best_evidence_page(pdf, evidence: str) -> tuple[int, float]:
    """Return the page whose text best covers the evidence bigrams."""
    best_page = 1
    best_score = 0.0
    for index, page in enumerate(pdf.pages, 1):
        page_text = norm_text(page.extract_text() or "")
        if not page_text:
            continue
        if evidence and evidence in page_text:
            return index, 1.0
        score = _evidence_containment(evidence, page_text)
        if score > best_score:
            best_page, best_score = index, score
    return best_page, best_score


def _validate_word_locators(file_path: Path, locators: dict[str, Any]) -> dict[str, Any]:
    article_no = str(locators.get("article_no") or "").strip()
    if not article_no or file_path.suffix.lower() != ".docx":
        return {"valid": True}
    try:
        document = Document(file_path)
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        if norm_text(article_no) not in norm_text(text):
            return {"valid": False, "errors": [f"Article '{article_no}' was not found in DOCX text"]}
        return {"valid": True}
    except Exception as exc:
        return {"valid": False, "errors": [f"Word validation error: {exc}"]}
